"""Document loading, chunking, and embedding.

Chunks carry true character spans into the source document. That is what lets an
evidence quote resolve to `document.raw_text[start:end]` — an exact highlight, not
"somewhere in this paragraph." Provenance is only as precise as the offsets you kept.
"""

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

import docx
import pypdf

from callosum.config import settings
from callosum.llm import embed  # re-exported: callers import it from here


@dataclass
class Chunk:
    ordinal: int
    text: str
    start_char: int  # inclusive offset into the document's raw_text
    end_char: int    # exclusive


def load(path: Path) -> str:
    """Extract plain text from a PDF, DOCX, or text file."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = pypdf.PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if suffix in (".txt", ".md", ".vtt"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {suffix}")


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# ~4 characters per token, the standard English approximation. Used only to *size*
# chunks, never to bill or budget. tiktoken would be wrong here anyway — it is
# OpenAI's tokenizer and mismatches both of our providers.
CHARS_PER_TOKEN = 4


def chunk(
    text: str, target_tokens: int | None = None, overlap_tokens: int | None = None
) -> list[Chunk]:
    """Split text into overlapping chunks on paragraph boundaries, preserving offsets.

    We pack whole paragraphs rather than slicing at a fixed offset. In a board
    transcript a paragraph is usually one speaker's turn, and cutting Priya's
    objection in half is how you lose the OPPOSED edge the system exists to record.
    Overlap carries trailing paragraphs forward, so a decision stated at a chunk
    boundary keeps its rationale attached.

    Each chunk is a *contiguous span* of the original text — carried-over paragraphs
    are adjacent to the ones that follow them — so `text[c.start_char:c.end_char]`
    reproduces the chunk exactly. Every downstream offset depends on that invariant.
    """
    cfg = settings()
    target = (target_tokens or cfg.chunk_tokens) * CHARS_PER_TOKEN
    overlap = (overlap_tokens or cfg.chunk_overlap_tokens) * CHARS_PER_TOKEN

    # (start, end) spans of each paragraph in the ORIGINAL text — not stripped copies,
    # so offsets stay true.
    spans = [
        (m.start(), m.end())
        for m in re.finditer(r"[^\n]+(?:\n(?!\s*\n)[^\n]+)*", text)
        if text[m.start() : m.end()].strip()
    ]
    if not spans:
        return []

    chunks: list[Chunk] = []
    current: list[tuple[int, int]] = []

    def flush() -> None:
        if not current:
            return
        start, end = current[0][0], current[-1][1]
        chunks.append(Chunk(len(chunks), text[start:end], start, end))

    def size_of(group: list[tuple[int, int]]) -> int:
        return group[-1][1] - group[0][0] if group else 0

    for span in spans:
        para_len = span[1] - span[0]

        # A single paragraph over budget (a wall-of-text PDF page) is split on
        # sentence boundaries. Offsets still come from the original text.
        if para_len > target:
            for m in re.finditer(r"[^.!?]+(?:[.!?]+|$)", text[span[0] : span[1]]):
                sub = (span[0] + m.start(), span[0] + m.end())
                if current and size_of(current) + (sub[1] - sub[0]) > target:
                    flush()
                    current = _carry_overlap(current, overlap)
                current.append(sub)
            continue

        if current and size_of(current) + para_len > target:
            flush()
            current = _carry_overlap(current, overlap)

        current.append(span)

    flush()
    return chunks


def _carry_overlap(
    current: list[tuple[int, int]], overlap: int
) -> list[tuple[int, int]]:
    """Keep trailing paragraphs from the previous chunk, up to the overlap budget."""
    carried: list[tuple[int, int]] = []
    size = 0
    for span in reversed(current):
        length = span[1] - span[0]
        if size + length > overlap:
            break
        carried.insert(0, span)
        size += length
    return carried


# Glyph equivalence classes: the same character in different encodings. Models
# habitually emit typographic punctuation ('smart quotes') where the source has
# ASCII, and vice versa. Treating those as mismatches would count faithful quotes
# as fabrications and corrupt the failure statistics. This is normalization, not
# fuzziness — a paraphrase still fails.
_GLYPH_CLASSES = {
    "'": "['‘’ʼ]",
    "‘": "['‘’ʼ]",
    "’": "['‘’ʼ]",
    '"': "[\"“”]",
    "“": "[\"“”]",
    "”": "[\"“”]",
    "-": "[-‐‑–—]",
    "–": "[-‐‑–—]",
    "—": "[-‐‑–—]",
    "…": r"(?:…|\.\.\.)",
}


def _token_pattern(token: str) -> str:
    return "".join(_GLYPH_CLASSES.get(ch, re.escape(ch)) for ch in token)


#: What may sit between two tokens of a located quote: horizontal whitespace, or a
#: single line break with indentation on either side — never a blank line.
#:
#: `\r` is in the horizontal class rather than treated as a line terminator so CRLF
#: input reflows like LF. Leaving it out silently broke every CRLF document, which is
#: most of what arrives from Windows and from email.
#:
#: The lookahead is `[ \t\r]*\n` rather than `\s*\n` so it cannot itself consume a
#: newline while deciding whether a newline follows.
_REFLOW = r"(?:[ \t\r]*\n[ \t\r]*(?![ \t\r]*\n)|[ \t\r]+)"


def locate(quote: str, haystack: str) -> tuple[int, int] | None:
    r"""Find a quote's exact character span, tolerating reflowed whitespace and
    typographic-vs-ASCII punctuation.

    Models reflow quotes across line breaks and swap quote/dash glyphs — the text is
    faithful but the bytes are not. So we match a whitespace-flexible, glyph-class
    regex, and return offsets into the *original* string so the span highlights
    correctly.

    Deliberately NOT fuzzy beyond whitespace, case, and glyph equivalence. A
    paraphrase is a fabrication, and it must not be located.

    Reflow is matched WITHIN a block, never ACROSS one. `\s+` would happily bridge a
    blank line, so a quote could be assembled from the end of one speaker's turn and
    the start of the next — text that appears verbatim in the document while never
    having been said by anyone. `_REFLOW` allows horizontal whitespace and at most one
    line break with surrounding indentation, and the negative lookahead refuses a
    second line break.

    THIS BOUNDS THE THESIS: an edge is verified because its evidence quote was located,
    so what `locate()` accepts is what "verified" means. Widening it is a claim about
    the corpus and must be measured, not argued.
    """
    tokens = quote.split()
    if not tokens:
        return None

    pattern = _REFLOW.join(_token_pattern(t) for t in tokens)
    match = re.search(pattern, haystack, flags=re.IGNORECASE)
    return (match.start(), match.end()) if match else None


__all__ = ["Chunk", "load", "content_hash", "chunk", "locate", "embed"]
