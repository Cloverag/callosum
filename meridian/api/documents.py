import uuid
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, UploadFile, status
from pydantic import BaseModel, Field

from meridian import documents as domain
from meridian.api.deps import CurrentPrincipal

router = APIRouter(prefix="/api/documents", tags=["documents"])


class IntakeDocumentRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200, description="Title of the board document")
    doc_type: str = Field("transcript", description="Document type (e.g. transcript, deck, memo, minutes)")
    raw_text: str = Field(..., min_length=1, description="Raw text or transcript body")
    sensitivity: int = Field(2, ge=0, le=4, description="Sensitivity clearance level (0=Public, 4=Restricted)")
    source_uri: str | None = Field(None, description="Source URI or file path reference")


@router.get("")
def list_documents(
    principal: CurrentPrincipal, doc_type: str | None = None
) -> list[domain.Document]:
    """Documents the caller may read, newest ingestion first."""
    return domain.list_documents(
        workspace_id=principal.workspace_id,
        clearance=principal.clearance,
        doc_type=doc_type,
    )


@router.get("/quarantine")
def list_quarantine(principal: CurrentPrincipal) -> list[domain.QuarantineItem]:
    """List extraction failures / quarantine items for the tenant workspace."""
    return domain.list_quarantine(workspace_id=principal.workspace_id)


@router.get("/{document_id}")
def get_document(document_id: uuid.UUID, principal: CurrentPrincipal) -> domain.Document:
    """One document."""
    return domain.get_document(
        str(document_id), workspace_id=principal.workspace_id, clearance=principal.clearance
    )


@router.post("/intake", status_code=status.HTTP_201_CREATED)
def intake_document(
    req: IntakeDocumentRequest, principal: CurrentPrincipal
) -> domain.Document:
    """Intake a plain text document into the tenant workspace (Meridian P4)."""
    return domain.intake_document(
        title=req.title,
        doc_type=req.doc_type,
        raw_text=req.raw_text,
        sensitivity=req.sensitivity,
        source_uri=req.source_uri,
        workspace_id=principal.workspace_id,
        author_principal_id=principal.id,
    )


@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_document(
    principal: CurrentPrincipal,
    file: UploadFile = File(...),
    title: str | None = Form(None),
    doc_type: str = Form("transcript"),
    sensitivity: int = Form(2),
) -> domain.Document:
    """Upload a file (PDF, DOCX, TXT, VTT, MD) into the tenant workspace (Meridian P4)."""
    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    filename = file.filename or "uploaded_document.txt"
    doc_title = title.strip() if title and title.strip() else filename

    # Extract text according to file extension
    try:
        if filename.lower().endswith(".pdf"):
            import pypdf
            import io
            reader = pypdf.PdfReader(io.BytesIO(content))
            raw_text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif filename.lower().endswith(".docx"):
            import docx
            import io
            d = docx.Document(io.BytesIO(content))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            raw_text = "\n".join(parts)
        else:
            raw_text = content.decode("utf-8", errors="replace")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse file content: {exc}",
        )

    if not raw_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File contains no extractable plain text",
        )

    try:
        return domain.intake_document(
            title=doc_title,
            doc_type=doc_type,
            raw_text=raw_text,
            sensitivity=sensitivity,
            source_uri=filename,
            workspace_id=principal.workspace_id,
            author_principal_id=principal.id,
        )
    except domain.DuplicateDocumentError as exc:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail=str(exc))
    except domain.InvalidSensitivityError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))
    except domain.DocumentError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))
